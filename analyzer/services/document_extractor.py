"""
Document Text Extractor Service
Extracts text from PDF, DOCX, and image files in memory.
Uses OpenAI Vision API or Tesseract OCR for image text extraction.
Implements clean separation of concerns with clear error handling.
"""
import io
import base64
import re
import logging
from PIL import Image, UnidentifiedImageError
import pdfplumber
from docx import Document
import openai
from django.conf import settings

# Try to import pytesseract for local OCR fallback
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentExtractionError(Exception):
    """Base exception for document extraction errors."""
    pass


class UnsupportedFileTypeError(DocumentExtractionError):
    """Raised when file type is not supported."""
    pass


class FileTooLargeError(DocumentExtractionError):
    """Raised when file exceeds size limit."""
    pass


class ExtractionFailedError(DocumentExtractionError):
    """Raised when text extraction fails."""
    pass


class InsufficientTextError(DocumentExtractionError):
    """Raised when extracted text is below minimum threshold."""
    pass


class FileTypeDetector:
    """
    Utility class for detecting and validating file types.
    """
    # Allowed extensions and their MIME categories
    ALLOWED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',  # Treat .doc as docx (python-docx handles both)
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image',
    }
    
    # Maximum file size: 2MB
    MAX_FILE_SIZE = 2 * 1024 * 1024
    
    @classmethod
    def detect(cls, filename: str) -> str:
        """
        Detect file type from filename extension.
        
        Args:
            filename: Name of the file
            
        Returns:
            File type category: 'pdf', 'docx', or 'image'
            
        Raises:
            UnsupportedFileTypeError: If extension is not allowed
        """
        ext = '.' + filename.lower().split('.')[-1] if '.' in filename else ''
        
        if ext not in cls.ALLOWED_EXTENSIONS:
            allowed = ', '.join(sorted(cls.ALLOWED_EXTENSIONS.keys()))
            raise UnsupportedFileTypeError(
                f"File type '{ext}' is not supported. "
                f"Please upload one of: {allowed}"
            )
        
        return cls.ALLOWED_EXTENSIONS[ext]
    
    @classmethod
    def validate_size(cls, file_obj) -> None:
        """
        Validate file size is within limits.
        
        Args:
            file_obj: Django UploadedFile object
            
        Raises:
            FileTooLargeError: If file exceeds max size
        """
        if file_obj.size > cls.MAX_FILE_SIZE:
            max_mb = cls.MAX_FILE_SIZE / (1024 * 1024)
            raise FileTooLargeError(
                f"File is too large. Maximum allowed size is {max_mb:.0f}MB."
            )
        
        if file_obj.size == 0:
            raise ExtractionFailedError("The uploaded file is empty.")


class PDFExtractor:
    """
    Extracts text from PDF files using pdfplumber.
    Processes entirely in memory.
    """
    
    @staticmethod
    def extract(file_content: bytes) -> str:
        """
        Extract text from PDF bytes.
        
        Args:
            file_content: Raw PDF file bytes
            
        Returns:
            Extracted text content
            
        Raises:
            ExtractionFailedError: If PDF cannot be parsed
        """
        try:
            pdf_file = io.BytesIO(file_content)
            text_parts = []
            
            with pdfplumber.open(pdf_file) as pdf:
                if len(pdf.pages) == 0:
                    raise ExtractionFailedError("PDF file contains no pages.")
                
                for i, page in enumerate(pdf.pages):
                    # Timeout protection: limit to first 20 pages
                    if i >= 20:
                        logger.warning("PDF has >20 pages, stopping extraction")
                        break
                    
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if not text_parts:
                raise ExtractionFailedError(
                    "Could not extract text from PDF. "
                    "The file may be image-based or scanned. "
                    "Try uploading an image screenshot instead."
                )
            
            return '\n\n'.join(text_parts)
            
        except ExtractionFailedError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ExtractionFailedError(
                f"Failed to read PDF file. It may be corrupted or password-protected."
            )


class DOCXExtractor:
    """
    Extracts text from DOCX files using python-docx.
    Processes entirely in memory.
    """
    
    @staticmethod
    def extract(file_content: bytes) -> str:
        """
        Extract text from DOCX bytes.
        
        Args:
            file_content: Raw DOCX file bytes
            
        Returns:
            Extracted text content
            
        Raises:
            ExtractionFailedError: If DOCX cannot be parsed
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if not text_parts:
                raise ExtractionFailedError(
                    "Could not extract text from document. "
                    "The file appears to be empty."
                )
            
            return '\n'.join(text_parts)
            
        except ExtractionFailedError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ExtractionFailedError(
                f"Failed to read document file. It may be corrupted or in an unsupported format."
            )


class ImageOCRExtractor:
    """
    Extracts text from images using OpenAI Vision API (primary) and Tesseract OCR (fallback).
    Processes entirely in memory.
    
    OpenAI is preferred for its accuracy with resume images, handwritten text, and scanned documents.
    Tesseract is available as an offline fallback if OpenAI API is not configured.
    """
    
    # Prompts for text extraction optimized for resume/JD analysis
    RESUME_EXTRACTION_PROMPT = """CRITICAL: Extract EVERY character visible in this resume image.
1. Extract all text: headings, names, dates, numbers, symbols
2. Include: EXPERIENCE, EDUCATION, SKILLS, CERTIFICATIONS sections
3. Include: job titles, companies, dates, achievements with metrics
4. Include: education details, schools, degrees, GPA
5. Include: all skills and certifications
6. Extract even faint or hard-to-read text
7. Preserve formatting: bullets, indentation, line breaks
8. Output ONLY raw text - no summaries or reorganization"""

    JOB_DESCRIPTION_EXTRACTION_PROMPT = """CRITICAL: Extract EVERY character visible in this job description image.
1. Extract all text: title, company, requirements, responsibilities
2. Include: all section headings (Requirements, Responsibilities, Benefits, Qualifications)
3. Include: every requirement and qualification
4. Include: all duties and responsibilities
5. Include: salary, benefits, location, work type info
6. Include: application instructions or contact info
7. Extract even faint or hard-to-read text
8. Preserve formatting: bullets, numbering, indentation, line breaks
9. Output ONLY raw text - no summaries or reorganization"""

    @classmethod
    def extract(cls, file_content: bytes, document_type: str = "jd") -> str:
        """
        Extract text from image bytes.
        Tries multiple methods in order:
        1. OpenAI Vision API (primary - cloud-based, most accurate for resumes)
        2. Tesseract OCR (fallback - local, free)
        
        Args:
            file_content: Raw image file bytes
            document_type: Type of document - 'resume' or 'jd' (job description). Defaults to 'jd'.
            
        Returns:
            Extracted text content
            
        Raises:
            ExtractionFailedError: If extraction fails with all methods
        """
        try:
            # Validate and load the image
            image_file = io.BytesIO(file_content)
            image = Image.open(image_file)
            image.load()
            
            # Convert to RGB if needed (handles RGBA, P mode, etc.)
            if image.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                if image.mode in ('RGBA', 'LA'):
                    background.paste(image, mask=image.split()[-1])
                else:
                    background.paste(image)
                image = background
            elif image.mode != 'RGB':
                image = image.convert('RGB')
            
        except UnidentifiedImageError:
            raise ExtractionFailedError(
                "Could not read this image file. It may be corrupted or in an unsupported format. "
                "Please upload a PNG/JPG/JPEG image."
            )
        except Exception as e:
            logger.error(f"Image loading error: {e}")
            raise ExtractionFailedError("Failed to read image file. Please ensure it's a valid image.")

        # Basic preprocessing to improve OCR reliability: autocontrast, contrast boost, sharpen
        try:
            processed_image = image.copy()
            from PIL import ImageEnhance, ImageFilter, ImageOps

            # Autocontrast to improve overall contrast
            processed_image = ImageOps.autocontrast(processed_image)

            # Increase contrast moderately
            enhancer = ImageEnhance.Contrast(processed_image)
            processed_image = enhancer.enhance(1.5)

            # Slight sharpening
            processed_image = processed_image.filter(ImageFilter.SHARPEN)
        except Exception:
            # If preprocessing fails, fall back to original image
            processed_image = image

        # Try Method 1: OpenAI Vision API (primary - cloud-based, most accurate)
        if settings.OPENAI_API_KEY:
            try:
                logger.info(f"Attempting image text extraction using OpenAI Vision API (document_type={document_type})")
                return cls._extract_with_openai(processed_image, document_type)
            except ExtractionFailedError:
                raise
            except Exception as e:
                logger.warning(f"OpenAI Vision extraction failed: {e}")
        else:
            logger.info("OpenAI API key not configured, attempting fallback methods")
        
        # Try Method 2: Tesseract OCR (fallback - local, free)
        if TESSERACT_AVAILABLE:
            try:
                logger.info("Attempting image text extraction using Tesseract OCR (fallback)")
                extracted_text = pytesseract.image_to_string(image).strip()
                
                if extracted_text:
                    logger.info(f"Successfully extracted {len(extracted_text)} chars from image via Tesseract")
                    return extracted_text
                else:
                    logger.info("Tesseract OCR returned empty text")
            except Exception as e:
                logger.warning(f"Tesseract OCR failed: {e}")
        else:
            logger.info("Tesseract OCR not available")
        
        # All methods failed - provide helpful error message
        error_msg = "Could not extract text from this image. "
        
        if not settings.OPENAI_API_KEY and not TESSERACT_AVAILABLE:
            error_msg += (
                "No OCR methods are configured. "
                "Please either: "
                "1) Set up OpenAI API key (see OPENAI_SETUP.md) - RECOMMENDED, or "
                "2) Install Tesseract OCR locally, or "
                "3) Paste the text manually."
            )
        elif not settings.OPENAI_API_KEY:
            error_msg += (
                "OpenAI API is not configured, and Tesseract extraction failed. "
                "Please set up OpenAI API (recommended), try a clearer image, or paste the text manually."
            )
        elif not TESSERACT_AVAILABLE:
            error_msg += (
                "OpenAI extraction failed, and Tesseract OCR is not available. "
                "Please try a clearer image, install Tesseract OCR, or paste the text manually."
            )
        else:
            error_msg += (
                "Please ensure the image contains clear, readable text, or paste the text manually."
            )
        
        raise ExtractionFailedError(error_msg)
    
    @classmethod
    def _extract_with_openai(cls, image: Image, document_type: str = "jd") -> str:
        """
        Extract text from PIL Image using OpenAI Vision API.
        Uses resume-specific or job-description-specific prompts for better accuracy.
        
        Args:
            image: PIL Image object (must be RGB)
            document_type: 'resume' or 'jd' to select appropriate extraction prompt
            
        Returns:
            Extracted text
            
        Raises:
            ExtractionFailedError: If OpenAI extraction fails
        """
        try:
            # Select appropriate prompt based on document type
            prompt = cls.RESUME_EXTRACTION_PROMPT if document_type == "resume" else cls.JOB_DESCRIPTION_EXTRACTION_PROMPT
            
            # Resize if too large (OpenAI has limits)
            max_dimension = 2048
            if max(image.size) > max_dimension:
                ratio = max_dimension / max(image.size)
                new_size = (int(image.size[0] * ratio), int(image.size[1] * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Convert to base64 with high quality for better OCR
            buffer = io.BytesIO()
            image.save(buffer, format='JPEG', quality=95, optimize=False)
            base64_image = base64.b64encode(buffer.getvalue()).decode('utf-8')
            
            # Call OpenAI Vision API
            client = openai.OpenAI(
                api_key=settings.OPENAI_API_KEY,
                base_url=settings.OPENAI_API_BASE,
            )
            
            response = client.chat.completions.create(
                model=settings.OPENAI_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=8000,
                temperature=0.0,
            )
            
            extracted_text = response.choices[0].message.content.strip()
            
            if not extracted_text:
                raise ExtractionFailedError(
                    "OpenAI could not detect any text in the image. "
                    "Please ensure the image contains readable text and try again."
                )
            
            logger.info(f"Successfully extracted {len(extracted_text)} chars from image via OpenAI")
            
            # If extraction seems too short, retry with more aggressive prompting
            if len(extracted_text) < 200:
                logger.warning(f"Extraction too short ({len(extracted_text)} chars), retrying with aggressive extraction...")
                try:
                    response_retry = client.chat.completions.create(
                        model=settings.OPENAI_MODEL,
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": prompt + "\n\n*** IMPORTANT: Be EXTREMELY thorough. Extract EVERYTHING visible, even small text. ***"
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/jpeg;base64,{base64_image}",
                                            "detail": "high"
                                        }
                                    }
                                ]
                            }
                        ],
                        max_tokens=8000,
                        temperature=0.0,
                    )
                    retry_text = response_retry.choices[0].message.content.strip()
                    if len(retry_text) > len(extracted_text):
                        logger.info(f"Retry successful: extracted {len(retry_text)} chars (improved from {len(extracted_text)})") 
                        extracted_text = retry_text
                except Exception as retry_err:
                    logger.warning(f"Retry failed: {retry_err}, using original extraction")
            
            return extracted_text
            
        except openai.APIError as e:
            logger.error(f"OpenAI API error during image extraction: {e}")
            raise ExtractionFailedError(
                "Failed to process image due to an API error. Please try again later."
            )
        except openai.AuthenticationError:
            logger.error("OpenAI authentication failed for image extraction")
            raise ExtractionFailedError(
                "Image text extraction is not properly configured."
            )
        except ExtractionFailedError:
            raise
        except Exception as e:
            logger.error(f"OpenAI image extraction error: {e}", exc_info=True)
            raise ExtractionFailedError(
                "Failed to process image. Please try a different image or paste the text manually."
            )


class DocumentTextExtractor:
    """
    Main service class for extracting text from uploaded documents.
    Provides a clean interface for the Django view layer.
    
    Usage:
        text = DocumentTextExtractor.extract(uploaded_file, filename)
    """
    
    # Minimum characters required for valid documents
    # Note: Images may extract less due to OCR limitations, so use a lower threshold
    MIN_TEXT_LENGTH = 200  # For PDFs and DOCX files
    MIN_TEXT_LENGTH_IMAGE = 100  # Lower threshold for images due to OCR variability
    
    @classmethod
    def extract(cls, file_obj, filename: str) -> str:
        """
        Extract and validate text from an uploaded file.
        
        Args:
            file_obj: Django UploadedFile object
            filename: Original filename
            
        Returns:
            Sanitized extracted text
            
        Raises:
            DocumentExtractionError subclass on failure
        """
        # Step 1: Validate file type
        file_type = FileTypeDetector.detect(filename)
        
        # Step 2: Validate file size
        FileTypeDetector.validate_size(file_obj)
        
        # Step 3: Read file content into memory
        file_content = file_obj.read()
        
        # Step 4: Extract text based on file type
        if file_type == 'pdf':
            raw_text = PDFExtractor.extract(file_content)
        elif file_type == 'docx':
            raw_text = DOCXExtractor.extract(file_content)
        elif file_type == 'image':
            raw_text = ImageOCRExtractor.extract(file_content)
            logger.debug(f"Image OCR returned {len(raw_text)} raw characters before sanitization")
        else:
            raise UnsupportedFileTypeError(f"Unknown file type: {file_type}")
        
        # Step 5: Sanitize text
        sanitized_text = cls._sanitize_text(raw_text)
        
        # Step 6: Validate minimum length
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        is_image = file_extension in ('png', 'jpg', 'jpeg', 'gif', 'bmp')
        min_length = cls.MIN_TEXT_LENGTH_IMAGE if is_image else cls.MIN_TEXT_LENGTH

        if len(sanitized_text) < min_length:
            if is_image:
                # For images, allow processing but log a clear warning so analysis proceeds
                logger.warning(
                    f"Image text extraction returned {len(sanitized_text)} characters "
                    f"(minimum expected: {min_length}). Analysis will proceed but may be less accurate."
                )
            else:
                raise InsufficientTextError(
                    f"Extracted text is too short ({len(sanitized_text)} characters). "
                    f"Job descriptions should be at least {cls.MIN_TEXT_LENGTH} characters. "
                    "Please upload a more complete job posting or paste the text manually."
                )
        
        logger.info(f"Successfully extracted {len(sanitized_text)} characters from {filename}")
        return sanitized_text
    
    @staticmethod
    def _sanitize_text(text: str) -> str:
        """
        Sanitize extracted text by removing artifacts and normalizing whitespace.
        Prepares text for safe LLM input.
        """
        # Remove null bytes and other control characters (except newlines/tabs)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize various unicode whitespace to standard space
        text = re.sub(r'[\u00a0\u2000-\u200b\u202f\u205f\u3000]', ' ', text)
        
        # Collapse multiple spaces into single space
        text = re.sub(r' +', ' ', text)
        
        # Collapse more than 2 consecutive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Final strip
        return text.strip()
