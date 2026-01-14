"""
CV Parser Service
Extracts text from PDF and DOCX files in memory without disk storage.
"""
import io
import logging
from typing import Optional
import pdfplumber
from docx import Document
from django.conf import settings

logger = logging.getLogger(__name__)


class CVParserService:
    """
    Service for parsing CV files (PDF and DOCX) in memory.
    All processing is done in-memory using BytesIO, no files are written to disk.
    """
    
    @staticmethod
    def validate_file(file_obj, filename: str) -> None:
        """
        Validate file size and extension.
        
        Args:
            file_obj: Django UploadedFile object
            filename: Name of the uploaded file
            
        Raises:
            ValueError: If validation fails
        """
        # Check file size
        if file_obj.size > settings.MAX_UPLOAD_SIZE:
            raise ValueError(
                f"File size exceeds maximum allowed size of "
                f"{settings.MAX_UPLOAD_SIZE / (1024 * 1024):.1f}MB"
            )
        
        # Check file extension
        extension = filename.lower().split('.')[-1]
        if f'.{extension}' not in settings.ALLOWED_FILE_EXTENSIONS:
            raise ValueError(
                f"File type '.{extension}' not supported. "
                f"Allowed types: {', '.join(settings.ALLOWED_FILE_EXTENSIONS)}"
            )
        
        # Check if file is empty
        if file_obj.size == 0:
            raise ValueError("Uploaded file is empty")
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file using pdfplumber.
        
        Args:
            file_content: Raw bytes of the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If PDF cannot be parsed or is empty
        """
        try:
            # Create in-memory file object
            pdf_file = io.BytesIO(file_content)
            
            text_content = []
            
            with pdfplumber.open(pdf_file) as pdf:
                if len(pdf.pages) == 0:
                    raise ValueError("PDF file contains no pages")
                
                # Extract text from each page
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            # Combine all pages
            full_text = '\n\n'.join(text_content).strip()
            
            if not full_text:
                raise ValueError("PDF file contains no readable text")
            
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            return full_text
            
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_content: Raw bytes of the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If DOCX cannot be parsed or is empty
        """
        try:
            # Create in-memory file object
            docx_file = io.BytesIO(file_content)
            
            # Parse DOCX document
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            # Combine all text
            full_text = '\n'.join(text_content).strip()
            
            if not full_text:
                raise ValueError("DOCX file contains no readable text")
            
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
    
    @classmethod
    def parse_cv(cls, file_obj, filename: str) -> str:
        """
        Main entry point for parsing CV files.
        Validates file and extracts text based on file type.
        
        Args:
            file_obj: Django UploadedFile object
            filename: Name of the uploaded file
            
        Returns:
            Extracted text content from the CV
            
        Raises:
            ValueError: If file validation or parsing fails
        """
        # Validate file
        cls.validate_file(file_obj, filename)
        
        # Read file content into memory
        file_content = file_obj.read()
        
        # Determine file type and parse accordingly
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            return cls.extract_text_from_pdf(file_content)
        elif extension == 'docx':
            return cls.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: .{extension}")
